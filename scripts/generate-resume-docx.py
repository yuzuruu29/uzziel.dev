"""Generate ATS-friendly resume DOCX from resume-content.ts data (parsed manually)."""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
TS_SOURCE = ROOT / "src" / "lib" / "resume-content.ts"
OUT_PORTFOLIO = ROOT / "public" / "downloads" / "uzziel-malolos-resume.docx"
OUT_DOWNLOADS = Path(r"C:\Users\uzzie\Downloads") / "Malolos Resume.docx"
OUT_TXT = ROOT / "public" / "downloads" / "uzziel-malolos-resume.txt"


def parse_ts_object(source: str, key: str) -> str:
    match = re.search(rf"{key}:\s*'((?:\\'|[^'])*)'", source)
    if not match:
        match = re.search(rf'{key}:\s*"([^"]*)"', source)
    if not match:
        raise ValueError(f"Missing key: {key}")
    return match.group(1).replace("\\'", "'")


def parse_string_array_block(source: str, key: str) -> list[str]:
    block = re.search(rf"{key}:\s*\[(.*?)\],", source, re.S)
    if not block:
        return []
    return [m.group(1).replace("\\'", "'") for m in re.finditer(r"'((?:\\'|[^'])*)'", block.group(1))]


def parse_skills(source: str) -> list[dict]:
    skills: list[dict] = []
    for group_match in re.finditer(
        r"group:\s*'((?:\\'|[^'])*)',\s*items:\s*\[(.*?)\]",
        source,
        re.S,
    ):
        group = group_match.group(1).replace("\\'", "'")
        items = [
            m.group(1).replace("\\'", "'")
            for m in re.finditer(r"'((?:\\'|[^'])*)'", group_match.group(2))
        ]
        skills.append({"group": group, "items": items})
    return skills


def parse_experience(source: str) -> list[dict]:
    entries: list[dict] = []
    for block in re.finditer(
        r"role:\s*'((?:\\'|[^'])*)'.*?org:\s*'((?:\\'|[^'])*)'.*?period:\s*'((?:\\'|[^'])*)'.*?bullets:\s*\[(.*?)\]",
        source,
        re.S,
    ):
        bullets = [
            m.group(1).replace("\\'", "'")
            for m in re.finditer(r"'((?:\\'|[^'])*)'", block.group(4))
        ]
        entries.append(
            {
                "role": block.group(1).replace("\\'", "'"),
                "org": block.group(2).replace("\\'", "'"),
                "period": block.group(3).replace("\\'", "'"),
                "bullets": bullets,
            }
        )
    return entries


def parse_projects(source: str) -> list[dict]:
    projects: list[dict] = []
    for block in re.finditer(
        r"title:\s*'((?:\\'|[^'])*)',\s*summary:\s*'((?:\\'|[^'])*)'",
        source,
        re.S,
    ):
        projects.append(
            {
                "title": block.group(1).replace("\\'", "'"),
                "summary": block.group(2).replace("\\'", "'"),
            }
        )
    return projects


def parse_education(source: str) -> list[dict]:
    entries: list[dict] = []
    for block in re.finditer(
        r"degree:\s*'((?:\\'|[^'])*)'.*?school:\s*'((?:\\'|[^'])*)'.*?year:\s*'((?:\\'|[^'])*)'.*?note:\s*'((?:\\'|[^'])*)'",
        source,
        re.S,
    ):
        entries.append(
            {
                "degree": block.group(1).replace("\\'", "'"),
                "school": block.group(2).replace("\\'", "'"),
                "year": block.group(3).replace("\\'", "'"),
                "note": block.group(4).replace("\\'", "'"),
            }
        )
    return entries


def load_resume_data() -> dict:
    source = TS_SOURCE.read_text(encoding="utf-8")
    return {
        "name": parse_ts_object(source, "name"),
        "tagline": parse_ts_object(source, "tagline"),
        "phone": parse_ts_object(source, "phone"),
        "email": parse_ts_object(source, "email"),
        "location": parse_ts_object(source, "location"),
        "linkedin": parse_ts_object(source, "linkedin"),
        "website": parse_ts_object(source, "website"),
        "summary": parse_ts_object(source, "summary"),
        "skills": parse_skills(source),
        "experience": parse_experience(source),
        "projects": parse_projects(source),
        "education": parse_education(source),
        "awards": parse_string_array_block(source, "awards"),
    }


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_body(doc: Document, text: str, bullet: bool = False) -> None:
    style = "List Bullet" if bullet else None
    p = doc.add_paragraph(text, style=style)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(2)


def build_docx(data: dict) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(data["name"])
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Calibri"

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tag.add_run(data["tagline"])
    tr.font.size = Pt(10.5)
    tr.font.name = "Calibri"

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = contact.add_run(
        f"{data['phone']}  •  {data['email']}  •  {data['location']}\n"
        f"{data['website'].replace('https://', '')}  •  LinkedIn"
    )
    cr.font.size = Pt(9.5)
    cr.font.name = "Calibri"

    add_heading(doc, "Professional Summary")
    add_body(doc, data["summary"])

    add_heading(doc, "Core Skills")
    for skill in data["skills"]:
        add_body(doc, f"{skill['group']}: {', '.join(skill['items'])}")

    add_heading(doc, "Relevant Experience")
    for job in data["experience"]:
        add_body(doc, f"{job['org']}")
        add_body(doc, f"{job['role']}  |  {job['period']}")
        for bullet in job["bullets"]:
            add_body(doc, bullet, bullet=True)

    add_heading(doc, "Selected Projects (2026)")
    for project in data["projects"]:
        add_body(doc, f"{project['title']} — {project['summary']}")

    add_heading(doc, "Education")
    for edu in data["education"]:
        add_body(doc, f"{edu['school']}  |  {edu['year']}")
        add_body(doc, f"{edu['degree']}  |  {edu['note']}")

    add_heading(doc, "Awards & Achievements")
    for award in data["awards"]:
        add_body(doc, award, bullet=True)

    add_body(doc, f"Portfolio and references: {data['website'].replace('https://', '')}")
    return doc


def build_txt(data: dict) -> str:
    lines = [
        data["name"],
        "",
        data["tagline"],
        f"{data['phone']} • {data['email']} • {data['location']}",
        f"{data['website'].replace('https://', '')} • LinkedIn",
        "",
        "PROFESSIONAL SUMMARY",
        "",
        data["summary"],
        "",
        "CORE SKILLS",
        "",
    ]
    for skill in data["skills"]:
        lines.append(f"{skill['group']}: {', '.join(skill['items'])}")
    lines.extend(["", "RELEVANT EXPERIENCE", ""])
    for job in data["experience"]:
        lines.extend([job["org"], f"{job['role']} | {job['period']}"])
        for bullet in job["bullets"]:
            lines.append(f"- {bullet}")
        lines.append("")
    lines.extend(["SELECTED PROJECTS (2026)", ""])
    for project in data["projects"]:
        lines.append(f"{project['title']} — {project['summary']}")
    lines.extend(["", "EDUCATION", ""])
    for edu in data["education"]:
        lines.extend(
            [
                f"{edu['school']} | {edu['year']}",
                f"{edu['degree']} | {edu['note']}",
                "",
            ]
        )
    lines.extend(["AWARDS & ACHIEVEMENTS", ""])
    for award in data["awards"]:
        lines.append(f"- {award}")
    lines.append("")
    lines.append(f"Portfolio and references: {data['website'].replace('https://', '')}")
    return "\n".join(lines)


def main() -> None:
    data = load_resume_data()
    OUT_PORTFOLIO.parent.mkdir(parents=True, exist_ok=True)
    doc = build_docx(data)
    doc.save(OUT_PORTFOLIO)
    doc.save(OUT_DOWNLOADS)
    OUT_TXT.write_text(build_txt(data), encoding="utf-8")
    print(json.dumps({"docx": str(OUT_PORTFOLIO), "downloads": str(OUT_DOWNLOADS), "txt": str(OUT_TXT)}, indent=2))


if __name__ == "__main__":
    main()